#!/usr/bin/env python3
"""Replace every Markdown file in ZIP archives with a readable DOCX file."""

from __future__ import annotations

import argparse
import importlib.util
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT_DIR = REPO_ROOT / "outputs" / "juristen-fallauswahl-2026-07-21"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "outputs" / "juristen-fallauswahl-2026-07-22-docx"
TABLE_HELPER = Path(
    "/Users/aaronsteiner/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.715.12143/skills/documents/scripts/table_geometry.py"
)
LINKED_MARKDOWN = re.compile(r"(?P<prefix>\]\()(?P<target>[^)\s]+)\.md(?P<suffix>(?:[?#][^)]*)?\))")


def load_table_helper():
    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_HELPER)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load table helper: {TABLE_HELPER}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


TABLE_GEOMETRY = load_table_helper()


def set_font(style, name: str, size: float, *, bold: bool = False, color: str = "000000") -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def set_paragraph_style(
    style,
    *,
    before: float,
    after: float,
    line_spacing: float,
    keep_with_next: bool = False,
) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def build_reference_docx(pandoc: str, output_path: Path) -> None:
    result = subprocess.run(
        [pandoc, "--print-default-data-file", "reference.docx"],
        check=True,
        stdout=subprocess.PIPE,
    )
    output_path.write_bytes(result.stdout)

    doc = Document(output_path)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_font(normal, "Calibri", 11)
    set_paragraph_style(normal, before=0, after=6, line_spacing=1.25)

    style_specs = {
        "Title": (20, "1F4D78", 0, 10),
        "Subtitle": (12, "555555", 0, 12),
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
        "Heading 4": (11, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in style_specs.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        set_font(style, "Calibri", size, bold=name.startswith("Heading"), color=color)
        set_paragraph_style(
            style,
            before=before,
            after=after,
            line_spacing=1.15,
            keep_with_next=True,
        )

    for name in ("Body Text", "First Paragraph", "Compact"):
        if name in doc.styles:
            set_font(doc.styles[name], "Calibri", 11)
            set_paragraph_style(doc.styles[name], before=0, after=6, line_spacing=1.25)

    for name in ("Footnote Text", "Endnote Text"):
        if name in doc.styles:
            set_font(doc.styles[name], "Calibri", 9)
            set_paragraph_style(doc.styles[name], before=0, after=2, line_spacing=1.1)

    for name in ("Caption", "Table Caption", "Image Caption"):
        if name in doc.styles:
            set_font(doc.styles[name], "Calibri", 9, color="555555")
            set_paragraph_style(doc.styles[name], before=4, after=4, line_spacing=1.1)

    if "Source Code" in doc.styles:
        set_font(doc.styles["Source Code"], "Courier New", 9, color="333333")
        set_paragraph_style(doc.styles["Source Code"], before=4, after=6, line_spacing=1.0)

    if "Table" not in doc.styles:
        doc.styles.add_style("Table", WD_STYLE_TYPE.TABLE)
    set_font(doc.styles["Table"], "Calibri", 9.5)

    doc.core_properties.title = "LAB-EU juristisches Prüfpaket"
    doc.core_properties.author = "LAB-EU"
    doc.core_properties.comments = "Automatisch aus Markdown für die juristische Durchsicht konvertiert."
    doc.save(output_path)


def rewrite_markdown_references(text: str) -> str:
    text = LINKED_MARKDOWN.sub(
        lambda match: f"{match.group('prefix')}{match.group('target')}.docx{match.group('suffix')}",
        text,
    )
    return re.sub(r"\.md\b", ".docx", text, flags=re.IGNORECASE)


def table_weights(table) -> list[float]:
    weights: list[float] = []
    for column in table.columns:
        lengths = [len(cell.text.strip()) for cell in column.cells]
        representative = max(lengths, default=1)
        weights.append(max(1.0, min(4.0, math.sqrt(max(1, representative)) / 3)))
    return weights


def postprocess_docx(path: Path) -> None:
    doc = Document(path)
    for table in doc.tables:
        widths = TABLE_GEOMETRY.column_widths_from_weights(table_weights(table), 9360)
        TABLE_GEOMETRY.apply_table_geometry(
            table,
            widths,
            table_width_dxa=9360,
            indent_dxa=120,
            cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
        )
        for row_index, row in enumerate(table.rows):
            row.height = None
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.1
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(9.5)
            if row_index == 0:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    doc.save(path)


def convert_markdown(pandoc: str, markdown_path: Path, archive_root: Path, reference: Path) -> None:
    source_text = markdown_path.read_text(encoding="utf-8", errors="replace")
    converted_text = rewrite_markdown_references(source_text)
    temporary_markdown = markdown_path.with_name(f".{markdown_path.name}.pandoc")
    temporary_markdown.write_text(converted_text, encoding="utf-8")
    output_path = markdown_path.with_suffix(".docx")
    try:
        result = subprocess.run(
            [
                pandoc,
                str(temporary_markdown),
                "--from=gfm+footnotes",
                "--to=docx",
                f"--reference-doc={reference}",
                f"--resource-path={markdown_path.parent}{os.pathsep}{archive_root}",
                "--wrap=none",
                f"--output={output_path}",
            ],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"Pandoc failed for {markdown_path}:\n{result.stdout}\n{result.stderr}"
            )
        postprocess_docx(output_path)
        if not output_path.is_file() or not zipfile.is_zipfile(output_path):
            raise RuntimeError(f"Invalid DOCX output: {output_path}")
        with zipfile.ZipFile(output_path) as zf:
            if "word/document.xml" not in zf.namelist():
                raise RuntimeError(f"DOCX has no word/document.xml: {output_path}")
        markdown_path.unlink()
    finally:
        temporary_markdown.unlink(missing_ok=True)


def create_zip(source_dir: Path, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(source_dir.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(source_dir))


def convert_archive(input_path: Path, output_path: Path, pandoc: str, reference: Path) -> dict:
    with tempfile.TemporaryDirectory(prefix="lab-eu-docx-") as temporary_directory:
        extraction_root = Path(temporary_directory) / "archive"
        extraction_root.mkdir()
        with zipfile.ZipFile(input_path) as zf:
            zf.extractall(extraction_root)

        markdown_files = sorted(extraction_root.rglob("*.md"))
        for index, markdown_path in enumerate(markdown_files, start=1):
            convert_markdown(pandoc, markdown_path, extraction_root, reference)
            if index % 50 == 0 or index == len(markdown_files):
                print(f"{input_path.name}: {index}/{len(markdown_files)} Markdown files converted")

        if list(extraction_root.rglob("*.md")):
            raise RuntimeError(f"Markdown files remain after conversion: {input_path}")
        docx_files = list(extraction_root.rglob("*.docx"))
        if len(docx_files) != len(markdown_files):
            raise RuntimeError(
                f"Expected {len(markdown_files)} DOCX files, found {len(docx_files)}"
            )
        create_zip(extraction_root, output_path)

    return {
        "input": str(input_path),
        "output": str(output_path),
        "converted_markdown_files": len(markdown_files),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-dir", type=Path, default=DEFAULT_INPUT_DIR)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    args = parser.parse_args()

    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise SystemExit("pandoc is required but was not found")

    input_dir = args.input_dir.resolve()
    output_dir = args.output_dir.resolve()
    inputs = sorted(input_dir.glob("*.zip"))
    if len(inputs) != 2:
        raise SystemExit(f"Expected exactly two ZIP files in {input_dir}, found {len(inputs)}")

    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="lab-eu-reference-") as temporary_directory:
        reference = Path(temporary_directory) / "reference.docx"
        build_reference_docx(pandoc, reference)
        results = [
            convert_archive(input_path, output_dir / input_path.name, pandoc, reference)
            for input_path in inputs
        ]
    print(json.dumps(results, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
