#!/usr/bin/env python3
"""Build professor-facing Word review packets from LAB-EU task/run artifacts."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
SKILL_ROOT = Path(
    "/Users/aaronsteiner/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.709.11516/skills/documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # type: ignore  # noqa: E402


OUTPUT_ROOT = REPO_ROOT / "outputs" / "professorenpakete_2026-07-11"
BASELINE_ROOT = (
    REPO_ROOT
    / "runs/baseline-deepseek-v4-pro-prof/20260711T075624Z/tasks"
)
HARNESS_ROOT = (
    REPO_ROOT
    / "runs/opencode-deepseek-v4-pro-prof/20260711T075626Z/tasks"
)

NAVY = "17365D"
BLUE = "2E74B5"
MID_BLUE = "DCE6F1"
LIGHT_BLUE = "EEF4FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE5"
TEXT = "1F2933"
MUTED = "5E6B78"
GREEN = "1B6B45"
GREEN_FILL = "EAF5EF"
RED = "9B1C1C"
RED_FILL = "FBECEC"
GOLD = "7A5A00"
GOLD_FILL = "FFF6D8"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}

FUNCTION_LABELS = {
    "structure": "Prüfungsaufbau",
    "legal_basis": "Rechtsgrundlage",
    "rule_statement": "Rechtssatz",
    "application": "Subsumtion",
    "argumentation": "Argumentation",
    "conclusion": "Ergebnis",
    "form_citation": "Form/Zitat",
}


@dataclass(frozen=True)
class CaseConfig:
    professor: str
    slug: str
    short_title: str
    task_dir: str
    field: str
    description: str
    filename: str
    editorial_notes: tuple[str, ...] = ()

    @property
    def task_path(self) -> Path:
        return REPO_ROOT / self.task_dir

    @property
    def run_key(self) -> str:
        return self.task_dir.removeprefix("tasks/").replace("/", "__")


CASES = (
    CaseConfig(
        "Krainer",
        "paypal",
        "PayPal-Käuferschutz schützt vor Klagen nicht",
        "tasks/de/zivilrecht/schuldrecht/fortgeschrittenenklausur-paypal-kauferschutz-schutzt-vor-klagen-nicht",
        "Zivilrecht · Schuldrecht",
        "Erfüllungswirkung einer PayPal-Zahlung, Käuferschutz, Versendungskauf, Verbrauchsgüterkauf und stellvertretendes Commodum nach § 285 BGB.",
        "01_PayPal_Kaeuferschutz_Pruefpaket.docx",
    ),
    CaseConfig(
        "Krainer",
        "ehe_espresso",
        "Ehe und Espresso",
        "tasks/de/zivilrecht/schuldrecht/examensklausur-ehe-und-espresso",
        "Zivilrecht · Schuld-/Sachenrecht",
        "Schlüsselgewalt, Eigentum und Mitbesitz an Haushaltsgegenständen, gutgläubiger Erwerb, § 1369 BGB und possessorischer Besitzschutz.",
        "02_Ehe_und_Espresso_Pruefpaket.docx",
    ),
    CaseConfig(
        "Krainer",
        "schiele",
        "Der gestohlene Schiele",
        "tasks/de/zivilrecht/sachenrecht/examensklausur-der-gestohlene-schiele",
        "Zivilrecht · Sachenrecht",
        "Eigentumsvermutung, Ersitzung, Erbengemeinschaft, Hinterlegung und gutgläubiger Auktionserwerb von Kulturgütern.",
        "03_Der_gestohlene_Schiele_Pruefpaket.docx",
        (
            "In der Word-Fassung wurde der offensichtliche Datumsfehler beim Ende der Ersitzungsfrist von 2005 auf 2015 berichtigt.",
        ),
    ),
    CaseConfig(
        "Ralf",
        "schmaehkritik",
        "Fall-05 Schmähkritik",
        "tasks/de/oeffentliches-recht/verwaltungsrecht/verpflichtungsklage/fall-05-verpflichtungsklage-auf-polizeiliches-einschreiten-wegen-schmaehkritik",
        "Verwaltungsrecht",
        "Verpflichtungsklage auf behördliches Einschreiten gegen ehrverletzende Onlineäußerungen, Medienaufsicht, Polizeirecht und Ermessensreduzierung.",
        "01_Fall_05_Schmaehkritik_Pruefpaket.docx",
        (
            "Die Beurteilung folgt dem im Fall bereitgestellten historischen RStV-/TMG- und Landesrechtsstand.",
        ),
    ),
    CaseConfig(
        "Ralf",
        "schredder",
        "40.000 t Schredder-Schrott",
        "tasks/de/oeffentliches-recht/verwaltungsrecht/examensklausur-40-000-t-schredder-schrott",
        "Verwaltungsrecht · Polizeirecht",
        "Eilrechtsschutz nach § 80 Abs. 5 VwGO, Dereliktion, Zustands- und Handlungsverantwortlichkeit, Zweckveranlassung und Störerauswahl.",
        "02_Schredder_Schrott_Pruefpaket.docx",
        (
            "Die Musterlösung enthält im Original einzelne redaktionelle Inkonsistenzen; in der Word-Fassung wurden nur eindeutige Norm- und Ergebnisverweise berichtigt.",
            "Der in der Musterlösung erwähnte Wohngebietsbezug steht nicht ausdrücklich im bereitgestellten Sachverhalt und bleibt als Prüfhinweis sichtbar.",
        ),
    ),
    CaseConfig(
        "Ralf",
        "atomkonsens",
        "Der Atomkonsens",
        "tasks/de/oeffentliches-recht/staatsrecht/ubungsfall-der-atomkonsens",
        "Staatsorganisationsrecht",
        "Bund-Länder-Streit, Bundesauftragsverwaltung nach Art. 85 GG, Sach- und Wahrnehmungskompetenz sowie informelles Verwaltungshandeln.",
        "03_Der_Atomkonsens_Pruefpaket.docx",
        (
            "Offensichtliche Extraktionsabbrüche und leere Überschriften der Musterlösung wurden in der Word-Fassung redaktionell geschlossen; die rechtliche Aussage wurde nicht verändert.",
        ),
    ),
    CaseConfig(
        "Svenja",
        "trio",
        "Ein zauderndes Trio",
        "tasks/de/strafrecht/materielles-strafrecht/ubungsfall-ein-zauderndes-trio",
        "Strafrecht · Allgemeiner Teil",
        "Versuchsbeginn bei Mittäterschaft, Lossagung eines Beteiligten sowie Rücktritt vom Versuch und von der Verbrechensverabredung.",
        "01_Ein_zauderndes_Trio_Pruefpaket.docx",
    ),
    CaseConfig(
        "Svenja",
        "kontaktlos",
        "Kontaktloses Bezahlen",
        "tasks/de/strafrecht/materielles-strafrecht/fortgeschrittenenklausur-kontaktloses-bezahlen",
        "Strafrecht · Besonderer Teil",
        "Kontaktloser Einsatz einer fremden EC-Karte mit Schwerpunkten bei Betrug, Computerbetrug und Urkundenunterdrückung.",
        "02_Kontaktloses_Bezahlen_Pruefpaket.docx",
    ),
    CaseConfig(
        "Svenja",
        "kunst_leben",
        "Kunst oder Leben",
        "tasks/de/oeffentliches-recht/staatsrecht/schwerpunktbereichsklausur-kunst-oder-leben",
        "Persönlichkeitsrecht · Grundrechte",
        "Kunstfreiheit und allgemeines Persönlichkeitsrecht bei Livestreams sowie Plattform-, Bild- und Wortberichterstattung.",
        "03_Kunst_oder_Leben_Pruefpaket.docx",
        (
            "Die Beurteilung der Plattformhaftung folgt dem im Originalfall verwendeten historischen TMG-Rechtsstand.",
        ),
    ),
)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_fill(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges: dict) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, text, fld_end])
    set_run_font(run, size=9, color=MUTED)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, color: str = BLUE):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_external_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def xml_safe(text: str) -> str:
    """Remove characters WordprocessingML cannot represent."""
    text = text.replace("Verm\x02gensvorteil", "Vermögensvorteil")
    return "".join(
        char
        for char in text
        if char in "\t\n\r" or ord(char) >= 0x20
    )


def style_document(doc: Document, running_title: str, professor: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.16

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 18, 9),
        ("Heading 2", 13.5, BLUE, 14, 6),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if style_name == "Heading 1":
            style.paragraph_format.page_break_before = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.16

    quote = styles["Quote"]
    quote.font.name = "Calibri"
    quote._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    quote._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    quote.font.size = Pt(9.2)
    quote.font.italic = False
    quote.font.color.rgb = rgb(MUTED)
    quote.paragraph_format.left_indent = Inches(0.28)
    quote.paragraph_format.right_indent = Inches(0.1)
    quote.paragraph_format.space_before = Pt(2)
    quote.paragraph_format.space_after = Pt(2)
    quote.paragraph_format.line_spacing = 1.05

    custom_specs = {
        "Legal Subheading": (11, NAVY, True, 7, 3),
        "Criterion Header K3": (10.5, NAVY, True, 8, 3),
        "Criterion Header K2": (10.5, TEXT, True, 8, 3),
        "Criterion Header K1": (10.5, MUTED, True, 8, 3),
        "Verdict Pass": (10.5, GREEN, True, 8, 3),
        "Verdict Fail": (10.5, RED, True, 8, 3),
        "Verdict Error": (10.5, GOLD, True, 8, 3),
        "Navigation Link": (11.5, BLUE, True, 3, 3),
        "Small Muted": (9, MUTED, False, 0, 3),
        "Section Lead": (11, TEXT, False, 0, 8),
        "Callout": (10.5, TEXT, False, 4, 5),
    }
    for name, (size, color, bold, before, after) in custom_specs.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12
        if "Header" in name or "Verdict" in name or name in {
            "Legal Subheading",
            "Navigation Link",
        }:
            style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"LAB-EU · Prüfpaket · {running_title}")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"Professor {professor}  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p)


def add_inline_markdown(
    paragraph,
    text: str,
    *,
    default_color: str = TEXT,
    default_size: float | None = 10.5,
) -> None:
    text = xml_safe(text)
    text = text.replace("\\_", "_")
    token_re = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*]+\*(?!\*)|(?<!\w)_[^_]+_(?!\w))"
    )
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                name="Courier New",
                size=9.5 if default_size is not None else None,
                color=default_color,
            )
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_external_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, color=default_color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size, color=default_color)


def clean_markdown(text: str, *, case_slug: str | None = None, gold: bool = False) -> str:
    text = xml_safe(text)
    text = text.replace("\r\n", "\n")
    text = re.sub(r"\u00ad", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    if gold and case_slug == "schiele":
        text = text.replace("30.9.2005 um 24:00 Uhr", "30.9.2015 um 24:00 Uhr")
    if gold and case_slug == "schredder":
        text = text.replace("§ 78 Abs. 2 Nr. 2 VwGO analog", "§ 78 Abs. 1 Nr. 2 VwGO analog")
        text = text.replace(
            "die aufschiebende Wirkung seines Widerspruchs deshalb wiederherstellen",
            "die aufschiebende Wirkung seiner Klage deshalb wiederherstellen",
        )
    if gold and case_slug == "atomkonsens":
        text = text.replace(
            "#### III. Antragsgegenstand, §§ 69, 64 Abs. 1 BVerfGG\n\nnisters ist geeignet",
            "#### III. Antragsgegenstand, §§ 69, 64 Abs. 1 BVerfGG\n\nDie Weisung des Bundesumweltministers ist geeignet",
        )
        text = text.replace(
            "#### I.\n\n93 Abs. 1 Nr. 3 GG i.V.m. §§ 13 Nr. 7, 68 ff. BVerfGG.",
            "#### I. Zuständigkeit des Bundesverfassungsgerichts\n\nDie Zuständigkeit des Bundesverfassungsgerichts ergibt sich aus Art. 93 Abs. 1 Nr. 3 GG i.V.m. §§ 13 Nr. 7, 68 ff. BVerfGG.",
        )
        text = text.replace(
            "#### VI.\n\n### B. Begründetheit",
            "#### VI. Zwischenergebnis\n\nDer Antrag ist zulässig.\n\n### B. Begründetheit",
        )
        text = text.replace("die Weisung, den D zu entlassen", "die Weisung, den D abzulösen")
        text = text.replace("BundLänder-Streit", "Bund-Länder-Streit")
    return text.strip()


def render_markdown(
    doc: Document,
    text: str,
    *,
    heading_map: dict[int, str] | None = None,
    skip_first_heading: bool = False,
) -> None:
    heading_map = heading_map or {
        1: "Heading 2",
        2: "Heading 3",
        3: "Legal Subheading",
        4: "Legal Subheading",
        5: "Legal Subheading",
        6: "Legal Subheading",
    }
    lines = text.splitlines()
    paragraph_lines: list[str] = []
    in_code = False
    code_lines: list[str] = []
    seen_heading = False

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        content = " ".join(line.strip() for line in paragraph_lines).strip()
        paragraph_lines = []
        if not content:
            return
        p = doc.add_paragraph()
        add_inline_markdown(p, content)

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                p = doc.add_paragraph(style="Quote")
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Courier New", size=8.8, color=TEXT)
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            title = re.sub(r"\*\*|__|`", "", heading_match.group(2)).strip()
            if skip_first_heading and not seen_heading:
                seen_heading = True
                continue
            seen_heading = True
            p = doc.add_paragraph(style=heading_map.get(level, "Legal Subheading"))
            add_inline_markdown(p, title, default_color=NAVY, default_size=None)
            continue
        if re.match(r"^\s*([-*_])(?:\s*\1){2,}\s*$", line):
            flush_paragraph()
            continue
        if not line.strip():
            flush_paragraph()
            continue
        blockquote = re.match(r"^>\s?(.*)$", line)
        if blockquote:
            flush_paragraph()
            p = doc.add_paragraph(style="Quote")
            add_inline_markdown(p, blockquote.group(1), default_color=MUTED)
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet.group(1))
            continue
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, numbered.group(1))
            continue
        paragraph_lines.append(line)
    flush_paragraph()


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
    p = doc.add_paragraph(style="Callout")
    set_paragraph_fill(p, fill)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=NAVY, bold=True)
    add_inline_markdown(p, text)


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cells[0], LIGHT_GRAY)
        for cell in cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": MID_GRAY},
                bottom={"val": "single", "sz": "4", "color": MID_GRAY},
                start={"val": "single", "sz": "4", "color": MID_GRAY},
                end={"val": "single", "sz": "4", "color": MID_GRAY},
            )
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9.5, color=TEXT)
    apply_table_geometry(
        table,
        [2200, 7160],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def format_rate(score: dict) -> str:
    return f"{score['n_passed']}/{score['n_criteria']} ({100 * score['criterion_pass_rate']:.1f} %)"


def criticality_entry(score: dict, tier: int) -> dict:
    prefix = f"{tier} ("
    for key, value in score.get("breakdown_by_criticality", {}).items():
        if key.startswith(prefix):
            return value
    return {"n_criteria": 0, "n_passed": 0, "pass_rate": 0.0}


def add_score_table(doc: Document, score: dict) -> None:
    k3 = criticality_entry(score, 3)
    k2 = criticality_entry(score, 2)
    k1 = criticality_entry(score, 1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    headers = ["Gesamt", "★★★", "★★", "★", "Fehler"]
    values = [
        format_rate(score),
        f"{k3['n_passed']}/{k3['n_criteria']} ({100 * k3.get('pass_rate', 0):.0f} %)",
        f"{k2['n_passed']}/{k2['n_criteria']} ({100 * k2.get('pass_rate', 0):.0f} %)",
        f"{k1['n_passed']}/{k1['n_criteria']} ({100 * k1.get('pass_rate', 0):.0f} %)",
        str(score["n_errors"]),
    ]
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(values):
        cell = table.rows[1].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9, color=TEXT, bold=(idx == 0))
    widths = [2200, 1800, 1800, 1800, 1760]
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": MID_GRAY},
                bottom={"val": "single", "sz": "4", "color": MID_GRAY},
                start={"val": "single", "sz": "4", "color": MID_GRAY},
                end={"val": "single", "sz": "4", "color": MID_GRAY},
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def flatten_outline(nodes: list[dict], depth: int = 1) -> list[tuple[dict, int]]:
    flattened: list[tuple[dict, int]] = []
    for node in nodes:
        flattened.append((node, depth))
        flattened.extend(flatten_outline(node.get("children", []), depth + 1))
    return flattened


def outline_groups(rubric: dict) -> list[tuple[dict, int, list[dict]]]:
    criteria = rubric["criteria"]
    by_outline: dict[str, list[dict]] = {}
    for criterion in criteria:
        outline_id = criterion.get("analysis_tags", {}).get("outline_id", "Ü")
        by_outline.setdefault(outline_id, []).append(criterion)
    groups: list[tuple[dict, int, list[dict]]] = []
    if by_outline.get("Ü"):
        groups.append(({"id": "Ü", "label": "Übergreifend"}, 1, by_outline.pop("Ü")))
    for node, depth in flatten_outline(rubric.get("outline", [])):
        groups.append((node, depth, by_outline.pop(node["id"], [])))
    for outline_id, unmatched in by_outline.items():
        groups.append(({"id": outline_id, "label": outline_id}, 1, unmatched))
    return groups


def add_group_heading(doc: Document, node: dict, depth: int) -> None:
    if depth == 1:
        style = "Heading 2"
    elif depth == 2:
        style = "Heading 3"
    else:
        style = "Legal Subheading"
    p = doc.add_paragraph(style=style)
    prefix = node.get("id", "")
    label = node.get("label", prefix)
    text = f"{prefix}  {label}" if prefix and prefix != label else label
    run = p.add_run(text)
    set_run_font(run, size=None, color=NAVY, bold=True)


def add_criterion_block(doc: Document, criterion: dict) -> None:
    tier = int(criterion.get("criticality", 2))
    style = f"Criterion Header K{tier}"
    fill = {3: MID_BLUE, 2: LIGHT_BLUE, 1: LIGHT_GRAY}[tier]
    p = doc.add_paragraph(style=style)
    set_paragraph_fill(p, fill)
    stars = "★" * tier
    function = FUNCTION_LABELS.get(
        criterion.get("analysis_tags", {}).get("function", ""),
        criterion.get("analysis_tags", {}).get("function", ""),
    )
    run = p.add_run(f"{criterion['id']} · {stars} · {criterion['title']}")
    set_run_font(run, size=10.5, color=NAVY if tier == 3 else TEXT, bold=True)
    if function:
        run = p.add_run(f"  [{function}]")
        set_run_font(run, size=9, color=MUTED, italic=True)

    p = doc.add_paragraph()
    run = p.add_run("Prüfmaßstab: ")
    set_run_font(run, size=10.2, color=NAVY, bold=True)
    add_inline_markdown(p, criterion["match_criteria"])
    notes = criterion.get("review_notes") or []
    if notes:
        p = doc.add_paragraph(style="Small Muted")
        run = p.add_run("Review-Hinweis: ")
        set_run_font(run, size=9, color=MUTED, bold=True)
        run = p.add_run(" ".join(notes))
        set_run_font(run, size=9, color=MUTED)


def add_rubric_section(doc: Document, rubric: dict) -> None:
    add_callout(
        doc,
        "Prüfauftrag",
        "Bitte kommentieren Sie Kriterien, die rechtlich falsch, unvollständig, redundant, falsch gewichtet oder gegenüber einer vertretbaren Gegenauffassung zu eng formuliert sind.",
    )
    for node, depth, criteria in outline_groups(rubric):
        if not criteria and depth > 1:
            continue
        add_group_heading(doc, node, depth)
        for criterion in criteria:
            add_criterion_block(doc, criterion)


def add_verdict_block(doc: Document, criterion: dict, result: dict) -> None:
    verdict = result.get("verdict", "error")
    tier = int(criterion.get("criticality", 2))
    function = FUNCTION_LABELS.get(
        criterion.get("analysis_tags", {}).get("function", ""),
        criterion.get("analysis_tags", {}).get("function", ""),
    )
    if verdict == "pass":
        icon, label, style, fill, color = "✓", "ERFÜLLT", "Verdict Pass", GREEN_FILL, GREEN
    elif verdict == "fail":
        icon, label, style, fill, color = "✗", "NICHT ERFÜLLT", "Verdict Fail", RED_FILL, RED
    else:
        icon, label, style, fill, color = "!", "FEHLER", "Verdict Error", GOLD_FILL, GOLD
    p = doc.add_paragraph(style=style)
    set_paragraph_fill(p, fill)
    run = p.add_run(
        f"{icon} {label} · {criterion['id']} · {'★' * tier} · {criterion['title']}"
    )
    set_run_font(run, size=10.5, color=color, bold=True)
    if function:
        run = p.add_run(f"  [{function}]")
        set_run_font(run, size=9, color=MUTED, italic=True)

    p = doc.add_paragraph()
    run = p.add_run("Rubrik: ")
    set_run_font(run, size=10.0, color=NAVY, bold=True)
    add_inline_markdown(
        p,
        criterion["match_criteria"],
        default_color=TEXT,
        default_size=10.0,
    )

    p = doc.add_paragraph()
    run = p.add_run("Judge-Begründung: ")
    set_run_font(run, size=10.2, color=NAVY, bold=True)
    run = p.add_run(xml_safe(result.get("reasoning", "Keine Begründung vorhanden.")))
    set_run_font(run, size=10.2, color=TEXT)
    evidence_items = result.get("evidence") or []
    for evidence in evidence_items:
        p = doc.add_paragraph(style="Quote")
        run = p.add_run(f"Fundstelle: „{xml_safe(evidence)}“")
        set_run_font(run, size=9.2, color=MUTED)
    if not evidence_items:
        p = doc.add_paragraph(style="Quote")
        run = p.add_run("Fundstelle: keine konkrete Textstelle angegeben")
        set_run_font(run, size=9.2, color=MUTED, italic=True)


def add_evaluation(doc: Document, rubric: dict, score: dict) -> None:
    add_score_table(doc, score)
    p = doc.add_paragraph(style="Small Muted")
    p.add_run(
        f"Judge: {score['judge_model']} · {score['votes_per_criterion']} Vote je Kriterium · "
        f"{score['n_errors']} technische Fehler"
    )
    results_by_id = {result["id"]: result for result in score["criteria_results"]}
    for node, depth, criteria in outline_groups(rubric):
        available = [criterion for criterion in criteria if criterion["id"] in results_by_id]
        if not available:
            continue
        add_group_heading(doc, node, depth)
        passed = sum(results_by_id[c["id"]].get("verdict") == "pass" for c in available)
        p = doc.add_paragraph(style="Small Muted")
        p.add_run(f"{passed}/{len(available)} Kriterien erfüllt")
        for criterion in available:
            add_verdict_block(doc, criterion, results_by_id[criterion["id"]])


def add_main_section(doc: Document, title: str, bookmark: str, bookmark_id: int, lead: str):
    p = doc.add_paragraph(style="Heading 1")
    add_bookmark(p, bookmark, bookmark_id)
    run = p.add_run(title)
    set_run_font(run, size=17, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Section Lead")
    p.add_run(lead)


def add_cover(
    doc: Document,
    config: CaseConfig,
    task: dict,
    rubric: dict,
    baseline_score: dict,
    harness_score: dict,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("JURISTISCHES PRÜFPAKET")
    set_run_font(run, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(config.short_title)
    set_run_font(run, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(17)
    run = p.add_run(config.description)
    set_run_font(run, size=12.5, color=MUTED)

    baseline_pct = 100 * baseline_score["criterion_pass_rate"]
    harness_pct = 100 * harness_score["criterion_pass_rate"]
    delta = harness_pct - baseline_pct
    add_metadata_table(
        doc,
        [
            ("Zuständig", f"Professor {config.professor}"),
            ("Rechtsgebiet", config.field),
            ("Rubrik", f"{len(rubric['criteria'])} Kriterien · GPT-5.5-generiert"),
            ("LLM-Baseline", f"{format_rate(baseline_score)}"),
            ("Agent-Harness", f"{format_rate(harness_score)}"),
            ("Differenz", f"{delta:+.1f} Prozentpunkte"),
            ("Judge", f"{baseline_score['judge_model']} · 1 Vote je Kriterium"),
        ],
    )

    add_callout(
        doc,
        "Ihr Auftrag",
        "Bitte prüfen Sie die juristische Qualität der Rubrik sowie die einzelnen Judge-Entscheidungen. Verwenden Sie dafür Word-Kommentare; die KI-Lösungen selbst sollen nicht redaktionell verbessert werden.",
    )
    if config.editorial_notes:
        add_callout(doc, "Redaktioneller Hinweis", " ".join(config.editorial_notes), GOLD_FILL)

    p = doc.add_paragraph(style="Small Muted")
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("ABSCHNITTSNAVIGATION")
    set_run_font(run, size=9, color=MUTED, bold=True)
    nav_items = (
        ("01  Fall: Aufgabe, Sachverhalt und Musterlösung", "sec_fall"),
        ("02  Rubriken", "sec_rubric"),
        ("03  LLM-Lösung (Baseline) mit Judge-Bewertung", "sec_baseline"),
        ("04  Agent-Lösung (Harness) mit Judge-Bewertung", "sec_harness"),
    )
    for label, anchor in nav_items:
        p = doc.add_paragraph(style="Navigation Link")
        add_internal_hyperlink(p, label, anchor)

    p = doc.add_paragraph(style="Small Muted")
    p.paragraph_format.space_before = Pt(10)
    p.add_run(
        "Navigation in Word: Ansicht → Navigationsbereich. Die Hauptabschnitte und Unterpunkte sind anklickbar und einklappbar."
    )


def document_files(task_path: Path) -> list[Path]:
    files = sorted(path for path in (task_path / "documents").rglob("*") if path.is_file())
    return files


def add_case_material(doc: Document, config: CaseConfig, task: dict) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Aufgabenbeschreibung")
    add_metadata_table(
        doc,
        [
            ("Titel", task.get("title", config.short_title)),
            ("Arbeitsauftrag", task.get("instructions", "")),
            ("Abgabe", str(task.get("deliverables", "fallloesung.md"))),
            ("Tags", ", ".join(task.get("tags", []))),
        ],
    )

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Fallbeschreibung und Anlagen")
    for source in document_files(config.task_path):
        p = doc.add_paragraph(style="Heading 3")
        p.add_run(source.name)
        try:
            text = source.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            p = doc.add_paragraph(style="Small Muted")
            p.add_run(f"Binäre Anlage: {source.name}")
            continue
        render_markdown(
            doc,
            clean_markdown(text, case_slug=config.slug),
            heading_map={1: "Legal Subheading", 2: "Legal Subheading", 3: "Legal Subheading", 4: "Legal Subheading", 5: "Legal Subheading", 6: "Legal Subheading"},
            skip_first_heading=True,
        )

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Musterlösung")
    solution = clean_markdown(
        (config.task_path / "evals" / "loesung.md").read_text(encoding="utf-8"),
        case_slug=config.slug,
        gold=True,
    )
    render_markdown(
        doc,
        solution,
        heading_map={1: "Heading 3", 2: "Heading 3", 3: "Legal Subheading", 4: "Legal Subheading", 5: "Legal Subheading", 6: "Legal Subheading"},
        skip_first_heading=True,
    )


def run_artifacts(config: CaseConfig, root: Path) -> tuple[Path, dict]:
    submission = root / config.run_key / "submission"
    solution = submission / "fallloesung.md"
    if not solution.exists():
        solution = submission / "fallloesung-sut.md"
    if not solution.exists():
        raise FileNotFoundError(f"Keine Falllösung im Run-Artefakt gefunden: {submission}")
    score = load_json(submission / "scores.json")
    return solution, score


def add_system_section(
    doc: Document,
    config: CaseConfig,
    rubric: dict,
    solution_path: Path,
    score: dict,
    label: str,
) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Lösung des Systems")
    add_callout(
        doc,
        "Hinweis",
        f"Die folgende {label} wird unverändert wiedergegeben. Bitte prüfen Sie anschließend, ob die Judge-Entscheidungen den jeweiligen Prüfmaßstab korrekt auf den Text anwenden.",
    )
    solution_text = clean_markdown(solution_path.read_text(encoding="utf-8"), case_slug=config.slug)
    render_markdown(
        doc,
        solution_text,
        heading_map={1: "Heading 3", 2: "Heading 3", 3: "Legal Subheading", 4: "Legal Subheading", 5: "Legal Subheading", 6: "Legal Subheading"},
        skip_first_heading=False,
    )
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Bewertung durch den Judge")
    add_evaluation(doc, rubric, score)


def build_packet(config: CaseConfig, output_path: Path) -> dict:
    task = load_json(config.task_path / "task.json")
    rubric = load_json(config.task_path / "evals" / "rubric.json")
    baseline_solution, baseline_score = run_artifacts(config, BASELINE_ROOT)
    harness_solution, harness_score = run_artifacts(config, HARNESS_ROOT)

    if baseline_score["n_criteria"] != len(rubric["criteria"]):
        raise ValueError(f"Baseline rubric mismatch: {config.short_title}")
    if harness_score["n_criteria"] != len(rubric["criteria"]):
        raise ValueError(f"Harness rubric mismatch: {config.short_title}")
    if baseline_score["n_errors"] or harness_score["n_errors"]:
        raise ValueError(f"Judge errors present: {config.short_title}")

    doc = Document()
    style_document(doc, config.short_title, config.professor)
    add_cover(doc, config, task, rubric, baseline_score, harness_score)

    add_main_section(
        doc,
        "01  Fall: Aufgabe, Sachverhalt und Musterlösung",
        "sec_fall",
        10,
        "Ausgangsmaterial für die juristische Prüfung. Maßgeblich ist der im Fall bereitgestellte historische Rechts- und Normenstand.",
    )
    add_case_material(doc, config, task)

    add_main_section(
        doc,
        "02  Rubriken",
        "sec_rubric",
        20,
        "GPT-5.5-generierte Kriterien, gegliedert nach Prüfungsstation und mit Wichtigkeitsstufe.",
    )
    add_rubric_section(doc, rubric)

    add_main_section(
        doc,
        "03  LLM-Lösung (Baseline) und Judge-Bewertung",
        "sec_baseline",
        30,
        "DeepSeek V4 Pro als einzelner Modellaufruf ohne Agent-Harness.",
    )
    add_system_section(
        doc,
        config,
        rubric,
        baseline_solution,
        baseline_score,
        "LLM-Baseline-Lösung",
    )

    add_main_section(
        doc,
        "04  Agent-Lösung (Harness) und Judge-Bewertung",
        "sec_harness",
        40,
        "DeepSeek V4 Pro im agentischen OpenCode-Harness mit derselben Rubrik und demselben Judge.",
    )
    add_system_section(
        doc,
        config,
        rubric,
        harness_solution,
        harness_score,
        "Agent-Harness-Lösung",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = f"Juristisches Prüfpaket – {config.short_title}"
    doc.core_properties.subject = "LAB-EU Rubrik- und Judge-Review"
    doc.core_properties.author = "LAB-EU"
    doc.core_properties.keywords = "LAB-EU, juristisches Review, Rubrik, Judge"
    doc.save(output_path)
    return {
        "title": config.short_title,
        "professor": config.professor,
        "criteria": len(rubric["criteria"]),
        "baseline": baseline_score["criterion_pass_rate"],
        "harness": harness_score["criterion_pass_rate"],
        "delta": harness_score["criterion_pass_rate"] - baseline_score["criterion_pass_rate"],
        "path": str(output_path),
    }


def build_instruction(output_path: Path) -> None:
    doc = Document()
    style_document(doc, "Kurzanleitung", "Review-Team")
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].font.color.rgb = rgb(TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("KURZANLEITUNG")
    set_run_font(run, size=10, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Konzept-Review der LAB-EU-Prüfpakete")
    set_run_font(run, size=23, color=TEXT, bold=True)
    p = doc.add_paragraph(style="Small Muted")
    p.add_run("Rubriken und ihre Anwendung durch den Review-Agenten · Stand 11. Juli 2026")

    p = doc.add_paragraph(style="Section Lead")
    run = p.add_run("Ziel: ")
    set_run_font(run, size=11, color=TEXT, bold=True)
    p.add_run(
        "Geprüft werden soll, ob die Rubrik juristisch fair, fallbezogen und sinnvoll gewichtet ist und ob der Review-Agent das jeweilige Kriterium anhand der betreffenden Lösung richtig bewertet hat. Es geht nicht darum, welches Lösungskonzept insgesamt besser abschneidet."
    )

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("1  Navigation")
    p = doc.add_paragraph()
    p.add_run(
        "In Word lässt sich über Ansicht → Navigationsbereich eine Registeransicht für Fall, Rubriken und die beiden Lösungsabschnitte öffnen. Die Überschriften sind anklickbar und einklappbar; auf der Titelseite führen zusätzliche Sprunglinks direkt zu den Abschnitten."
    )

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("2  Folgende Dinge gerne prüfen")
    for label, text in (
        (
            "Rubriken",
            "Sind die Kriterien rechtlich richtig, fallrelevant, klar prüfbar, vollständig und sinnvoll gewichtet?",
        ),
        (
            "Review-Agent",
            "Passen ERFÜLLT/NICHT ERFÜLLT, Begründung und Fundstellen zur jeweiligen Rubrik und zur Lösung?",
        ),
    ):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=10.5, color=TEXT, bold=True)
        p.add_run(text)

    p = doc.add_paragraph()
    p.add_run(
        "Hinweise können sich auf einzelne Kriterien oder allgemein auf die Rubrik und ihre Anwendung beziehen. Allgemeines Feedback ist ausdrücklich hilfreich, etwa wenn das Sternesystem grundsätzlich überzeugt, die aktuelle Einordnung aber noch nicht passt. In diesem Fall kann gerne beschrieben werden, wie die Grenzen zwischen drei, zwei und einem Stern gezogen werden sollten."
    )

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("3  Sternesystem")
    add_metadata_table(
        doc,
        [
            ("★★★", "ergebnistragend"),
            ("★★", "wichtig"),
            ("★", "Detail oder Ergänzung"),
        ],
    )
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("4  Feedback geben")
    p = doc.add_paragraph()
    p.add_run(
        "Für Feedback kann gerne die Kommentarfunktion in Word genutzt oder Text direkt im Dokument ergänzt werden. Eine eigene Rückgabedatei oder eine farbliche Codierung ist nicht nötig. Normale Markierungen sind selbstverständlich möglich."
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Kurzanleitung – Konzept-Review der LAB-EU-Prüfpakete"
    doc.core_properties.subject = "Anleitung für Professorinnen und Professoren"
    doc.core_properties.author = "LAB-EU"
    doc.save(output_path)


def build_archives(output_root: Path, instruction: Path, packet_paths: dict[str, list[Path]]) -> None:
    for professor, paths in packet_paths.items():
        archive = output_root / f"Professor_{professor}_Pruefpaket.zip"
        with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(instruction, arcname=instruction.name)
            for path in paths:
                zf.write(path, arcname=path.name)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_ROOT)
    parser.add_argument("--case", choices=[case.slug for case in CASES])
    parser.add_argument("--clean", action="store_true")
    args = parser.parse_args()

    output_root = args.output_dir.resolve()
    if args.clean and output_root.exists():
        shutil.rmtree(output_root)
    output_root.mkdir(parents=True, exist_ok=True)

    instruction = output_root / "00_Kurzanleitung_Professorenreview.docx"
    build_instruction(instruction)

    selected = [case for case in CASES if not args.case or case.slug == args.case]
    summaries = []
    packet_paths: dict[str, list[Path]] = {}
    for config in selected:
        professor_dir = output_root / f"Professor_{config.professor}"
        output_path = professor_dir / config.filename
        summaries.append(build_packet(config, output_path))
        packet_paths.setdefault(config.professor, []).append(output_path)

    if not args.case:
        build_archives(output_root, instruction, packet_paths)
    (output_root / "manifest.json").write_text(
        json.dumps(summaries, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(f"Wrote {len(summaries)} review packet(s) to {output_root}")


if __name__ == "__main__":
    main()
