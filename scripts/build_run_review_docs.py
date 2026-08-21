#!/usr/bin/env python3
"""Ein Word-Dokument je Lauf-Konfiguration ueber alle Faelle eines Tasksets.

Anders als build_professor_review_packets.py (ein Paket je Fall, zwei Systeme)
schneidet dieses Skript entlang der Konfiguration: eine Datei je Modell-Harness-
Kombination, darin alle Faelle mit Aufgabe, Sachverhalt, Musterloesung, der
KI-Loesung und der ausgewerteten Rubrik. Judge-Begruendungen erscheinen nur bei
nicht erfuellten Kriterien.
"""
from __future__ import annotations

import argparse
import json
import pathlib
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO = pathlib.Path(__file__).resolve().parents[1]
GEBIET = {"oeffentliches-recht": "Öffentliches Recht", "strafrecht": "Strafrecht", "zivilrecht": "Zivilrecht"}
GRUEN, ROT, GRAU = RGBColor(0x1B, 0x6B, 0x45), RGBColor(0x9B, 0x1C, 0x1C), RGBColor(0x5E, 0x6B, 0x78)
FUNKTION = {"structure": "Prüfungsaufbau", "legal_basis": "Rechtsgrundlage", "rule_statement": "Rechtssatz",
            "application": "Subsumtion", "argumentation": "Argumentation", "conclusion": "Ergebnis",
            "form_citation": "Form/Zitat"}


def md_to_doc(doc: Document, text: str, *, max_chars: int = 200_000) -> None:
    """Markdown grob nach Word: Ueberschriften, Listen, Absaetze."""
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Gekürzt]\n"
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("#"):
            lvl = len(line) - len(line.lstrip("#"))
            doc.add_paragraph(line.lstrip("# ").strip(), style=f"Heading {min(lvl + 2, 5)}")
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        elif re.match(r"^\s*\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\s*\d+\.\s", "", line), style="List Number")
        elif line.lstrip().startswith(">"):
            doc.add_paragraph(line.lstrip("> ").strip(), style="Quote")
        else:
            doc.add_paragraph(re.sub(r"\*\*(.+?)\*\*", r"\1", line))


def read_dir(d: pathlib.Path, suffixes=(".md",)) -> str:
    if not d.exists():
        return ""
    out = []
    for p in sorted(d.rglob("*")):
        if p.is_file() and p.suffix in suffixes:
            out.append(f"## {p.name}\n\n" + p.read_text(encoding="utf-8", errors="replace"))
    return "\n\n".join(out)


def case_sort_key(entry: dict) -> tuple:
    order = {"oeffentliches-recht": 0, "strafrecht": 1, "zivilrecht": 2}
    return (order.get(entry["gebiet_slug"], 9), entry["titel"].lower())


def collect(run_dir: pathlib.Path) -> list[dict]:
    out = []
    for scores_path in sorted(run_dir.glob("tasks/*/scores.json")):
        s = json.loads(scores_path.read_text(encoding="utf-8"))
        rubric_path = pathlib.Path(s["rubric"])
        task_dir = rubric_path.parent.parent
        if not task_dir.exists():
            print(f"  Ueberspringe {scores_path.parent.name}: Task fehlt", file=sys.stderr)
            continue
        task = json.loads((task_dir / "task.json").read_text(encoding="utf-8"))
        rubric = json.loads(rubric_path.read_text(encoding="utf-8"))
        rel = str(task_dir).split("/tasks/de/")[-1]
        sub = scores_path.parent / "submission"
        loesung = read_dir(sub) if sub.exists() else ""
        out.append({
            "titel": task.get("title") or task_dir.name,
            "gebiet_slug": rel.split("/")[0],
            "task_dir": task_dir,
            "task": task,
            "rubric": rubric,
            "scores": s,
            "loesung": loesung,
        })
    return sorted(out, key=case_sort_key)


def build(run_dir: pathlib.Path, titel: str, untertitel: str, ziel: pathlib.Path,
          *, mit_stil: bool = False) -> None:
    faelle = collect(run_dir)
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = "Calibri", Pt(10.5)

    doc.add_paragraph(titel, style="Title")
    p = doc.add_paragraph(untertitel); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ges_p = sum(f["scores"]["n_passed"] for f in faelle)
    ges_k = sum(f["scores"]["n_criteria"] for f in faelle)
    p = doc.add_paragraph(f"{len(faelle)} Fälle · {ges_k} Kriterien · {ges_p} erfüllt ({ges_p/ges_k*100:.1f} %)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Judge-Begründungen sind nur bei nicht erfüllten Kriterien abgedruckt.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_paragraph("Übersicht", style="Heading 1")
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Nr.", "Fall", "Rechtsgebiet", "Erfüllt"]):
        tbl.rows[0].cells[i].text = h
    for i, f in enumerate(faelle, 1):
        c = tbl.add_row().cells
        sc = f["scores"]
        c[0].text, c[1].text = str(i), f["titel"][:70]
        c[2].text = GEBIET.get(f["gebiet_slug"], f["gebiet_slug"])
        c[3].text = f"{sc['n_passed']}/{sc['n_criteria']} ({sc['n_passed']/sc['n_criteria']*100:.0f} %)"
    doc.add_page_break()

    for i, f in enumerate(faelle, 1):
        sc, task = f["scores"], f["task"]
        doc.add_paragraph(f"{i:02d}  {f['titel']}", style="Heading 1")
        meta = doc.add_paragraph()
        r = meta.add_run(f"{GEBIET.get(f['gebiet_slug'], f['gebiet_slug'])} · "
                         f"{sc['n_passed']}/{sc['n_criteria']} Kriterien erfüllt "
                         f"({sc['n_passed']/sc['n_criteria']*100:.1f} %)")
        r.font.color.rgb, r.font.size = GRAU, Pt(9)

        doc.add_paragraph("Aufgabenstellung", style="Heading 2")
        md_to_doc(doc, task.get("instructions", ""))
        doc.add_paragraph("Sachverhalt und Anlagen", style="Heading 2")
        md_to_doc(doc, read_dir(f["task_dir"] / "documents"))
        doc.add_paragraph("Musterlösung", style="Heading 2")
        md_to_doc(doc, read_dir(f["task_dir"] / "evals"))
        doc.add_paragraph("KI-Lösung", style="Heading 2")
        md_to_doc(doc, f["loesung"] or "[Keine Lösung abgelegt]")

        doc.add_paragraph("Bewertung", style="Heading 2")
        krit = {c["id"]: c for c in f["rubric"]["criteria"]}
        for res in sc["criteria_results"]:
            c = krit.get(res["id"], {})
            erfuellt = res["verdict"] == "pass"
            p = doc.add_paragraph()
            r = p.add_run(f"{'✓' if erfuellt else '✗'}  {res['id']}  {res.get('title','')}")
            r.bold, r.font.color.rgb = True, (GRUEN if erfuellt else ROT)
            fn = str((c.get("analysis_tags") or {}).get("function") or "")
            if fn:
                r2 = p.add_run(f"   [{FUNKTION.get(fn, fn)}]"); r2.font.color.rgb, r2.font.size = GRAU, Pt(8)
            if not erfuellt:
                if c.get("match_criteria"):
                    q = doc.add_paragraph(c["match_criteria"], style="Quote")
                    q.runs[0].font.size = Pt(9)
                begr = res.get("reasoning") or ""
                if not begr:
                    for v in res.get("votes", []):
                        if v.get("verdict") != "pass" and v.get("reasoning"):
                            begr = v["reasoning"]; break
                if begr:
                    pb = doc.add_paragraph()
                    rb = pb.add_run(f"Judge: {begr}")
                    rb.font.size, rb.font.color.rgb = Pt(9), GRAU

        stil = sc.get("style_score") or {}
        if mit_stil and stil.get("n_eligible"):
            doc.add_paragraph("Stilbewertung", style="Heading 2")
            p = doc.add_paragraph()
            r = p.add_run(f"{stil['n_passed']}/{stil['n_eligible']} stilfähige Kriterien erfüllt "
                          f"({stil['n_passed']/stil['n_eligible']*100:.1f} %). Bewertet werden nur "
                          f"Kriterien der Kategorien Subsumtion und Argumentation.")
            r.font.color.rgb, r.font.size = GRAU, Pt(9)
            for res in sc.get("style_results", []):
                if res.get("verdict") == "pass":
                    continue
                p = doc.add_paragraph()
                r = p.add_run(f"✗  {res['id']}  {res.get('title','')}")
                r.bold, r.font.color.rgb = True, ROT
                if res.get("reasoning"):
                    pb = doc.add_paragraph()
                    rb = pb.add_run(f"Judge: {res['reasoning']}")
                    rb.font.size, rb.font.color.rgb = Pt(9), GRAU
        doc.add_page_break()

    ziel.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = titel
    doc.core_properties.subject = "LAB-EU Laufauswertung"
    doc.core_properties.author = "LAB-EU"
    doc.save(ziel)
    print(f"  {ziel.name}: {len(faelle)} Fälle, {ges_k} Kriterien -> {ziel.stat().st_size/1e6:.1f} MB")



def fall_block(doc: Document, f: dict, konfigs: list[tuple[str, dict]], *, mit_stil: bool) -> None:
    """Fallmaterial einmal, danach je Konfiguration Loesung und Bewertung."""
    doc.add_paragraph("Aufgabenstellung", style="Heading 2")
    md_to_doc(doc, f["task"].get("instructions", ""))
    doc.add_paragraph("Sachverhalt und Anlagen", style="Heading 2")
    md_to_doc(doc, read_dir(f["task_dir"] / "documents"))
    doc.add_paragraph("Musterlösung", style="Heading 2")
    md_to_doc(doc, read_dir(f["task_dir"] / "evals"))

    krit = {c["id"]: c for c in f["rubric"]["criteria"]}
    for label, eintrag in konfigs:
        sc, loesung = eintrag["scores"], eintrag["loesung"]
        doc.add_page_break()
        doc.add_paragraph(label, style="Heading 1")
        p = doc.add_paragraph()
        r = p.add_run(f"{sc['n_passed']}/{sc['n_criteria']} Kriterien erfüllt "
                      f"({sc['n_passed']/sc['n_criteria']*100:.1f} %)")
        r.font.color.rgb, r.font.size = GRAU, Pt(9)
        doc.add_paragraph("KI-Lösung", style="Heading 2")
        md_to_doc(doc, loesung or "[Keine Lösung abgelegt]")
        doc.add_paragraph("Bewertung", style="Heading 2")
        for res in sc["criteria_results"]:
            c = krit.get(res["id"], {})
            erf = res["verdict"] == "pass"
            p = doc.add_paragraph()
            r = p.add_run(f"{'✓' if erf else '✗'}  {res['id']}  {res.get('title','')}")
            r.bold, r.font.color.rgb = True, (GRUEN if erf else ROT)
            fn = str((c.get("analysis_tags") or {}).get("function") or "")
            if fn:
                r2 = p.add_run(f"   [{FUNKTION.get(fn, fn)}]")
                r2.font.color.rgb, r2.font.size = GRAU, Pt(8)
            if not erf:
                if c.get("match_criteria"):
                    q = doc.add_paragraph(c["match_criteria"], style="Quote")
                    q.runs[0].font.size = Pt(9)
                begr = res.get("reasoning") or ""
                if not begr:
                    for v in res.get("votes", []):
                        if v.get("verdict") != "pass" and v.get("reasoning"):
                            begr = v["reasoning"]; break
                if begr:
                    pb = doc.add_paragraph()
                    rb = pb.add_run(f"Judge: {begr}")
                    rb.font.size, rb.font.color.rgb = Pt(9), GRAU
        stil = sc.get("style_score") or {}
        if mit_stil and stil.get("n_eligible"):
            doc.add_paragraph("Stilbewertung", style="Heading 2")
            p = doc.add_paragraph()
            r = p.add_run(f"{stil['n_passed']}/{stil['n_eligible']} stilfähige Kriterien erfüllt "
                          f"({stil['n_passed']/stil['n_eligible']*100:.1f} %)")
            r.font.color.rgb, r.font.size = GRAU, Pt(9)
            for res in sc.get("style_results", []):
                if res.get("verdict") == "pass":
                    continue
                p = doc.add_paragraph()
                r = p.add_run(f"✗  {res['id']}  {res.get('title','')}")
                r.bold, r.font.color.rgb = True, ROT
                if res.get("reasoning"):
                    pb = doc.add_paragraph()
                    rb = pb.add_run(f"Judge: {res['reasoning']}")
                    rb.font.size, rb.font.color.rgb = Pt(9), GRAU


def build_per_case(runs: list[tuple[str, pathlib.Path]], ziel_dir: pathlib.Path,
                   *, mit_stil: bool = False) -> None:
    """Ein Dokument je Fall, darin alle uebergebenen Konfigurationen."""
    gesammelt = [(label, {e["task_dir"]: e for e in collect(d)}) for label, d in runs]
    basis = gesammelt[0][1]
    ziel_dir.mkdir(parents=True, exist_ok=True)
    for i, (task_dir, f) in enumerate(sorted(basis.items(), key=lambda kv: case_sort_key(kv[1])), 1):
        konfigs = [(label, m[task_dir]) for label, m in gesammelt if task_dir in m]
        if len(konfigs) != len(gesammelt):
            print(f"  {f['titel'][:40]}: nur {len(konfigs)} von {len(gesammelt)} Konfigurationen", file=sys.stderr)
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name, st.font.size = "Calibri", Pt(10.5)
        doc.add_paragraph(f["titel"], style="Title")
        p = doc.add_paragraph(GEBIET.get(f["gebiet_slug"], f["gebiet_slug"]))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(" · ".join(
            f"{lb}: {e['scores']['n_passed']}/{e['scores']['n_criteria']}"
            f" ({e['scores']['n_passed']/e['scores']['n_criteria']*100:.0f} %)" for lb, e in konfigs))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph("Judge-Begründungen sind nur bei nicht erfüllten Kriterien abgedruckt.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        fall_block(doc, f, konfigs, mit_stil=mit_stil)
        name = task_dir.name
        for a, b in (("ä", "ae"), ("ö", "oe"), ("ü", "ue"), ("Ä", "Ae"), ("Ö", "Oe"),
                     ("Ü", "Ue"), ("ß", "ss")):
            name = name.replace(a, b)
        name = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")
        out = ziel_dir / f"{i:02d}_{name}.docx"
        doc.core_properties.title = f["titel"]
        doc.core_properties.subject = "LAB-EU Fallpaket"
        doc.core_properties.author = "LAB-EU"
        doc.save(out)
        print(f"  {out.name}: {len(konfigs)} Konfigurationen, {f['scores']['n_criteria']} Kriterien")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--run-dir", type=pathlib.Path, action="append", required=True,
                    help="Lauf-Verzeichnis; mehrfach fuer den Modus je Fall.")
    ap.add_argument("--label", action="append", default=[],
                    help="Anzeigename je --run-dir (Modus je Fall).")
    ap.add_argument("--je-fall", action="store_true",
                    help="Ein Dokument je Fall statt je Konfiguration.")
    ap.add_argument("--titel", default="")
    ap.add_argument("--untertitel", default="")
    ap.add_argument("--output", type=pathlib.Path, required=True)
    ap.add_argument("--mit-stil", action="store_true",
                    help="Stilbewertung je Fall ergaenzen (nur Laeufe mit style_score).")
    a = ap.parse_args()
    if a.je_fall:
        labels = a.label or [d.parent.name for d in a.run_dir]
        if len(labels) != len(a.run_dir):
            raise SystemExit("Zahl der --label muss zur Zahl der --run-dir passen.")
        build_per_case(list(zip(labels, a.run_dir)), a.output, mit_stil=a.mit_stil)
    else:
        build(a.run_dir[0], a.titel, a.untertitel, a.output, mit_stil=a.mit_stil)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
